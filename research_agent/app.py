"""
Multi-Agent Research System — Streamlit Dashboard
Usage: streamlit run app.py
"""

import streamlit as st
import asyncio
import threading
import queue
import sys
import io
import os
import re
import time
from pathlib import Path
from datetime import datetime

# ─── Path setup ───────────────────────────────────────────────
ROOT = Path(__file__).parent
sys.path.insert(0, str(ROOT))

st.set_page_config(
    page_title="Research Agent",
    page_icon="🔬",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ─── Stage definitions ────────────────────────────────────────
STAGES = [
    (1, "①", "Input",       "주제 파싱 & 전략 수립"),
    (2, "②", "Literature",  "자료 수집 & 요약"),
    (3, "③", "Ideation",    "아이디어 생성 & 평가"),
    (4, "④", "Execution",   "실험 설계 & 코드"),
    (5, "⑤", "Review",      "다각도 비판 검토"),
    (6, "⑥", "Output",      "최종 문서 작성"),
]
STAGE_MARKER_MAP = {"①": 1, "②": 2, "③": 3, "④": 4, "⑤": 5, "⑥": 6}

# ─── Session state defaults ───────────────────────────────────
_DEFAULTS = {
    "running": False,
    "done": False,
    "error": None,
    "logs": [],
    "current_stage": 0,
    "stage_logs": {i: [] for i in range(1, 7)},
    "stage_done": set(),
    "review_scores": [],
    "final_score": None,
    "result_md": None,
    "output_path": None,
    "topic_val": "",
    "mode_val": "academic",
    "elapsed": 0,
    "start_time": None,
    "view_file": None,   # 이전 결과 보기: 선택된 파일 경로
    "view_md": None,     # 이전 결과 보기: 로드된 내용
}
for _k, _v in _DEFAULTS.items():
    if _k not in st.session_state:
        st.session_state[_k] = _v


# ─── Stdout capture ───────────────────────────────────────────
class _QueueWriter:
    def __init__(self, q):
        self._q = q
        self._orig = sys.__stdout__

    def write(self, text):
        if text and text.strip():
            self._q.put(("log", text.strip()))
        try:
            self._orig and self._orig.write(text)
        except Exception:
            pass

    def flush(self):
        try:
            self._orig and self._orig.flush()
        except Exception:
            pass


# ─── Pipeline thread ──────────────────────────────────────────
def _run_pipeline(mode, topic, max_papers, max_revisions, verbose, log_q, result_q):
    from core.pipeline import ResearchPipeline
    old = sys.stdout
    sys.stdout = _QueueWriter(log_q)
    try:
        pipeline = ResearchPipeline(
            mode=mode, topic=topic,
            max_papers=max_papers, max_revisions=max_revisions,
            verbose=verbose,
        )
        result = asyncio.run(pipeline.run())
        result_q.put({"ok": True, "result": result})
    except Exception as e:
        result_q.put({"ok": False, "error": str(e)})
    finally:
        sys.stdout = old
        log_q.put(("done", None))


# ─── Past-file helpers ────────────────────────────────────────
def _list_past_files():
    """outputs/papers + outputs/reports의 .md 파일 목록 (최신순)"""
    files = []
    for subdir in ["papers", "reports"]:
        d = ROOT / "outputs" / subdir
        if d.exists():
            for f in sorted(d.glob("*.md"), key=lambda x: x.stat().st_mtime, reverse=True):
                files.append((subdir, f))
    return files  # [(subdir_str, Path), ...]

def _friendly_name(subdir, path):
    """academic_ViT_vs_CNN_20260628_205317.md  →  [📄] ViT vs CNN (2026-06-28 20:53)"""
    stem = path.stem
    m = re.search(r"_(\d{8})_(\d{6})$", stem)
    if m:
        d, t = m.group(1), m.group(2)
        ts   = f"{d[:4]}-{d[4:6]}-{d[6:8]} {t[:2]}:{t[2:4]}"
        stem = stem[: m.start()]
    else:
        ts = ""
    stem  = re.sub(r"^(academic|strategy)_", "", stem)
    label = stem.replace("_", " ")
    badge = "📄" if subdir == "papers" else "📊"
    suffix = f"  ({ts})" if ts else ""
    return f"{badge} {label}{suffix}"

# ─── Log parsers ──────────────────────────────────────────────
def _detect_stage(text):
    for marker, num in STAGE_MARKER_MAP.items():
        if marker in text:
            return num
    return None

def _parse_review_score(text):
    m = re.search(r"앙상블 점수:\s*([\d.]+)", text)
    return float(m.group(1)) if m else None


# ─── Export: Word ─────────────────────────────────────────────
def _to_docx(md_text, title):
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    doc.styles["Normal"].font.name = "맑은 고딕"

    h = doc.add_heading(title, 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(f"생성: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    code_block = False
    for line in md_text.split("\n"):
        if line.startswith("```"):
            code_block = not code_block
            continue
        if code_block:
            doc.add_paragraph(line).style = "No Spacing"
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), 1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), 2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), 3)
        elif line.startswith("#### "):
            doc.add_heading(line[5:].strip(), 4)
        elif re.match(r"^-{3,}$", line.strip()):
            doc.add_paragraph("─" * 40)
        elif line.strip():
            p = doc.add_paragraph()
            for part in re.split(r"(\*\*.*?\*\*|\*.*?\*)", line):
                if part.startswith("**") and part.endswith("**"):
                    p.add_run(part[2:-2]).bold = True
                elif part.startswith("*") and part.endswith("*"):
                    p.add_run(part[1:-1]).italic = True
                else:
                    p.add_run(re.sub(r"`([^`]+)`", r"\1", part))
        else:
            doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ─── Export: PDF ──────────────────────────────────────────────
def _to_pdf(md_text, title):
    from fpdf import FPDF

    class PDF(FPDF):
        def header(self):
            self.set_font_size(8)
            self.set_text_color(150)
            self.cell(0, 8, "Multi-Agent Research System", align="R")
            self.ln(2)
        def footer(self):
            self.set_y(-12)
            self.set_font_size(8)
            self.set_text_color(150)
            self.cell(0, 8, f"- {self.page_no()} -", align="C")

    pdf = PDF(orientation="P", unit="mm", format="A4")
    pdf.set_auto_page_break(auto=True, margin=18)
    pdf.set_margins(20, 20, 20)

    font_name = "Helvetica"
    for fp in [r"C:\Windows\Fonts\malgun.ttf", r"C:\Windows\Fonts\NanumGothic.ttf"]:
        if os.path.exists(fp):
            try:
                pdf.add_font("KR", fname=fp)
                font_name = "KR"
            except Exception:
                pass
            break

    def sf(size=10, bold=False):
        pdf.set_font(font_name, style="B" if bold and font_name == "Helvetica" else "", size=size)

    def cell(txt):
        try:
            pdf.multi_cell(0, 6, txt)
        except Exception:
            pdf.multi_cell(0, 6, txt.encode("latin-1", errors="replace").decode("latin-1"))

    pdf.add_page()
    sf(15, bold=True); pdf.set_text_color(20, 80, 160)
    pdf.multi_cell(0, 10, title, align="C")
    sf(8); pdf.set_text_color(120)
    pdf.cell(0, 6, datetime.now().strftime("%Y-%m-%d %H:%M"), align="C")
    pdf.ln(8); pdf.set_text_color(0)

    code_block = False
    for line in md_text.split("\n"):
        if line.startswith("```"):
            code_block = not code_block; continue
        clean = re.sub(r"\*\*(.*?)\*\*|\*(.*?)\*|`([^`]+)`",
                       lambda m: m.group(1) or m.group(2) or m.group(3), line)
        if code_block:
            sf(8); pdf.set_text_color(60); cell(line or " "); pdf.set_text_color(0); continue
        if line.startswith("# "):
            pdf.ln(4); sf(13, bold=True); pdf.set_text_color(20, 80, 160); cell(line[2:].strip()); pdf.set_text_color(0)
        elif line.startswith("## "):
            pdf.ln(3); sf(11, bold=True); pdf.set_text_color(40, 100, 170); cell(line[3:].strip()); pdf.set_text_color(0)
        elif line.startswith("### "):
            pdf.ln(2); sf(10, bold=True); pdf.set_text_color(60, 120, 180); cell(line[4:].strip()); pdf.set_text_color(0)
        elif re.match(r"^-{3,}$", line.strip()):
            pdf.ln(1); pdf.set_draw_color(200); pdf.line(20, pdf.get_y(), 190, pdf.get_y()); pdf.ln(3)
        elif clean.strip():
            sf(10); cell(clean)
        else:
            pdf.ln(3)

    return bytes(pdf.output())


# ─── CSS ──────────────────────────────────────────────────────
st.markdown("""<style>
.stage-card { border-radius:10px; padding:12px 6px; text-align:center; margin:2px; }
.st-wait  { background:#1e293b; border:1px solid #334155; }
.st-run   { background:#422006; border:1px solid #f59e0b; animation:pulse 1.5s infinite; }
.st-done  { background:#052e16; border:1px solid #16a34a; }
@keyframes pulse { 0%,100%{opacity:1} 50%{opacity:.7} }
.s-icon  { font-size:22px; }
.s-name  { font-size:11px; font-weight:700; color:#e2e8f0; margin-top:4px; }
.s-desc  { font-size:9px;  color:#94a3b8; margin-top:2px; }
.s-stat  { font-size:10px; color:#64748b; margin-top:4px; }
.stage-detail { background:#0f172a; border:1px solid #1e293b; border-radius:8px;
                padding:12px 16px; margin:4px 0; font-size:13px; color:#cbd5e1; }
.score-bar-wrap { background:#1e293b; border-radius:8px; padding:16px; margin:8px 0; }
</style>""", unsafe_allow_html=True)

# ─── Sidebar ──────────────────────────────────────────────────
with st.sidebar:
    st.markdown("## 🔬 Research Agent")
    st.divider()

    mode = st.selectbox(
        "📌 연구 모드",
        ["academic", "strategy"],
        format_func=lambda x: "📄 학술 논문 연구" if x == "academic" else "📊 AI 전략 리포트",
        disabled=st.session_state.running,
    )
    topic = st.text_area(
        "🔍 연구 주제",
        placeholder="예: ViT vs CNN image classification survey",
        height=100,
        disabled=st.session_state.running,
    )

    with st.expander("⚙️ 고급 설정"):
        max_papers    = st.slider("최대 논문/자료 수",   3, 20, 10, disabled=st.session_state.running)
        max_revisions = st.slider("최대 Revision 횟수", 1,  5,  2, disabled=st.session_state.running)
        verbose       = st.checkbox("상세 로그", value=False, disabled=st.session_state.running)

    st.divider()
    run_btn = st.button(
        "⏳ 연구 진행 중..." if st.session_state.running else "🚀 연구 시작",
        disabled=st.session_state.running or not topic.strip(),
        use_container_width=True, type="primary",
    )

    if not st.session_state.running and (st.session_state.done or st.session_state.error):
        if st.button("🔄 초기화", use_container_width=True):
            for k in list(st.session_state.keys()):
                del st.session_state[k]
            st.rerun()

    # 사이드바 진행 상태
    if st.session_state.running:
        cur = st.session_state.current_stage
        if cur > 0:
            _, circle, name, _ = STAGES[cur - 1]
            st.info(f"{circle} **{name}** 진행 중...")
        if st.session_state.start_time:
            elapsed = int(time.time() - st.session_state.start_time)
            st.caption(f"⏱ 경과: {elapsed // 60}분 {elapsed % 60}초")

    if st.session_state.final_score is not None:
        score = st.session_state.final_score
        st.divider()
        st.markdown("**📊 최종 리뷰 점수**")
        color = "#22c55e" if score >= 0.7 else "#f59e0b" if score >= 0.5 else "#ef4444"
        st.markdown(
            f'<div style="font-size:28px;font-weight:bold;color:{color};text-align:center">'
            f'{score:.2f}<span style="font-size:14px;color:#94a3b8"> / 1.0</span></div>',
            unsafe_allow_html=True,
        )
        st.progress(score)

    # ─── 이전 결과 보기 ──────────────────────────────────────
    st.divider()
    st.markdown("### 📁 이전 결과 보기")
    past = _list_past_files()
    if not past:
        st.caption("저장된 결과 없음")
    else:
        # 표시 이름 목록
        options = [("(선택하세요)", None)] + [
            (_friendly_name(sub, p), p) for sub, p in past
        ]
        labels = [o[0] for o in options]
        paths  = [o[1] for o in options]

        # 현재 선택 인덱스 복원
        cur_path = st.session_state.view_file
        try:
            cur_idx = paths.index(cur_path) if cur_path in paths else 0
        except ValueError:
            cur_idx = 0

        sel_idx = st.selectbox(
            "파일 선택",
            range(len(labels)),
            index=cur_idx,
            format_func=lambda i: labels[i],
            label_visibility="collapsed",
            disabled=st.session_state.running,
            key="past_file_select",
        )
        sel_path = paths[sel_idx]

        if sel_path and sel_path != st.session_state.view_file:
            st.session_state.view_file = sel_path
            st.session_state.view_md   = sel_path.read_text(encoding="utf-8")
            st.rerun()
        elif sel_path is None and st.session_state.view_file:
            st.session_state.view_file = None
            st.session_state.view_md   = None
            st.rerun()

        if st.session_state.view_file:
            if st.button("✖ 닫기", use_container_width=True):
                st.session_state.view_file = None
                st.session_state.view_md   = None
                st.rerun()


# ─── Start pipeline ───────────────────────────────────────────
if run_btn and topic.strip() and not st.session_state.running:
    lq = queue.Queue()
    rq = queue.Queue()
    st.session_state._log_q    = lq
    st.session_state._result_q = rq
    st.session_state.running      = True
    st.session_state.done         = False
    st.session_state.error        = None
    st.session_state.logs         = []
    st.session_state.current_stage = 0
    st.session_state.stage_logs   = {i: [] for i in range(1, 7)}
    st.session_state.stage_done   = set()
    st.session_state.review_scores = []
    st.session_state.final_score  = None
    st.session_state.result_md    = None
    st.session_state.output_path  = None
    st.session_state.topic_val    = topic.strip()
    st.session_state.mode_val     = mode
    st.session_state.start_time   = time.time()
    st.session_state.view_file    = None
    st.session_state.view_md      = None

    threading.Thread(
        target=_run_pipeline,
        args=(mode, topic.strip(), max_papers, max_revisions, verbose, lq, rq),
        daemon=True,
    ).start()
    st.rerun()


# ─── Poll queue ───────────────────────────────────────────────
_log_q    = st.session_state.get("_log_q")
_result_q = st.session_state.get("_result_q")

if st.session_state.running and _log_q:
    for _ in range(80):
        if _log_q.empty():
            break
        kind, text = _log_q.get_nowait()
        if kind == "done":
            st.session_state.running = False
            if _result_q and not _result_q.empty():
                res = _result_q.get_nowait()
                if res.get("ok"):
                    path = res["result"]["output_path"]
                    st.session_state.output_path = path
                    st.session_state.result_md   = Path(path).read_text(encoding="utf-8")
                    st.session_state.done        = True
                    st.session_state.current_stage = 7   # all done
                    for i in range(1, 7):
                        st.session_state.stage_done.add(i)
                    if st.session_state.review_scores:
                        st.session_state.final_score = st.session_state.review_scores[-1]
                else:
                    st.session_state.error = res.get("error", "알 수 없는 오류")
            break
        elif kind == "log":
            st.session_state.logs.append(text)
            stage = _detect_stage(text)
            if stage:
                prev = st.session_state.current_stage
                if prev and prev != stage:
                    st.session_state.stage_done.add(prev)
                st.session_state.current_stage = stage
            cur = st.session_state.current_stage
            if cur:
                st.session_state.stage_logs[cur].append(text)
            score = _parse_review_score(text)
            if score is not None:
                st.session_state.review_scores.append(score)


# ═══════════════════════════════════════════════════════════════
# MAIN UI
# ═══════════════════════════════════════════════════════════════
st.title("🔬 Multi-Agent Research System")

_no_active = (
    not st.session_state.running
    and not st.session_state.done
    and not st.session_state.error
)
if _no_active and not st.session_state.view_md:
    st.info("👈 사이드바에서 연구 모드와 주제를 입력한 후 **연구 시작** 버튼을 클릭하세요.")
    st.stop()

# 이전 결과 보기 모드 (진행/완료 결과 없을 때만 메인에 표시)
if _no_active and st.session_state.view_md:
    _v_md   = st.session_state.view_md
    _v_path = st.session_state.view_file
    _v_name = _v_path.stem if _v_path else "결과"
    _v_ts   = datetime.now().strftime("%Y%m%d_%H%M%S")

    st.info(f"📂 **이전 결과 보기** — `{_v_path.name if _v_path else ''}`")

    # 다운로드 버튼
    d1, d2, d3 = st.columns(3)
    with d1:
        st.download_button(
            "📄 Markdown 다운로드",
            data=_v_md.encode("utf-8"),
            file_name=f"{_v_name}.md",
            mime="text/markdown",
            use_container_width=True,
        )
    with d2:
        try:
            st.download_button(
                "📝 Word (.docx) 다운로드",
                data=_to_docx(_v_md, _v_name),
                file_name=f"{_v_name}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
        except Exception as e:
            st.warning(f"Word 생성 실패: {e}")
    with d3:
        try:
            st.download_button(
                "🖨️ PDF 다운로드",
                data=_to_pdf(_v_md, _v_name),
                file_name=f"{_v_name}.pdf",
                mime="application/pdf",
                use_container_width=True,
            )
        except Exception as e:
            st.warning(f"PDF 생성 실패: {e}")

    st.divider()
    tab1, tab2 = st.tabs(["📖 렌더링 보기", "📋 Markdown 원문"])
    with tab1:
        st.markdown(_v_md)
    with tab2:
        st.text_area("원문", _v_md, height=600, label_visibility="collapsed")
    st.stop()

# ─── 1. Stage cards ───────────────────────────────────────────
st.subheader("📊 진행 단계")
cols = st.columns(6)
cur  = st.session_state.current_stage
done_set = st.session_state.stage_done

for i, (num, circle, name, desc) in enumerate(STAGES):
    with cols[i]:
        if num in done_set or (st.session_state.done and num <= 6):
            cls, icon, stxt = "st-done", "✅", "완료"
        elif num == cur and st.session_state.running:
            cls, icon, stxt = "st-run",  "🔄", "진행 중"
        else:
            cls, icon, stxt = "st-wait", "⏳", "대기"

        st.markdown(
            f'<div class="stage-card {cls}">'
            f'<div class="s-icon">{icon}</div>'
            f'<div class="s-name">{circle} {name}</div>'
            f'<div class="s-desc">{desc}</div>'
            f'<div class="s-stat">{stxt}</div></div>',
            unsafe_allow_html=True,
        )

st.divider()

# ─── 2. Stage detail panels ───────────────────────────────────
st.subheader("📋 단계별 상세")

STAGE_ICONS = ["📋", "📚", "💡", "⚙️", "🔍", "📝"]

for num, circle, name, desc in STAGES:
    s_logs = st.session_state.stage_logs.get(num, [])
    is_done    = num in done_set or (st.session_state.done and num <= 6)
    is_running = (num == cur and st.session_state.running)
    is_waiting = not is_done and not is_running

    icon = "✅" if is_done else ("🔄" if is_running else "⏳")
    label = f"{icon} {circle} {name}  —  {desc}"

    with st.expander(label, expanded=(is_running or (is_done and num == 5))):
        if is_waiting and not s_logs:
            st.caption("대기 중...")
        elif s_logs:
            # 주요 정보 추출
            key_lines = [l for l in s_logs if not l.startswith("─") and len(l) > 5]

            # Stage 5 review score display
            if num == 5 and st.session_state.review_scores:
                scores = st.session_state.review_scores
                sc_cols = st.columns(len(scores))
                for idx, sc in enumerate(scores):
                    with sc_cols[idx]:
                        color = "🟢" if sc >= 0.7 else ("🟡" if sc >= 0.5 else "🔴")
                        st.metric(f"리뷰 #{idx+1}", f"{sc:.2f}", delta=None)
                if len(scores) > 1:
                    st.progress(scores[-1], text=f"최종 점수: {scores[-1]:.2f} / 1.0")
                else:
                    st.progress(scores[-1])
                st.divider()

            # Show key log lines
            filtered = []
            for l in key_lines:
                if any(skip in l for skip in ["──", "WARNING", "JSON 파싱"]):
                    continue
                filtered.append(l)
            if filtered:
                st.code("\n".join(filtered[-30:]), language=None)
        if is_running:
            st.markdown("🔄 *처리 중...*")

st.divider()

# ─── 3. Review score gauge (prominent) ────────────────────────
if st.session_state.review_scores:
    scores = st.session_state.review_scores
    final  = scores[-1]
    color  = "#22c55e" if final >= 0.7 else "#f59e0b" if final >= 0.5 else "#ef4444"
    label  = "우수 ✅" if final >= 0.7 else ("통과 ⚠️" if final >= 0.5 else "미흡 ❌")

    st.subheader("📊 리뷰 점수")
    mc, pc = st.columns([1, 3])
    with mc:
        st.markdown(
            f'<div style="background:#0f172a;border:1px solid #334155;border-radius:12px;'
            f'padding:20px;text-align:center">'
            f'<div style="font-size:13px;color:#94a3b8;margin-bottom:6px">최종 점수</div>'
            f'<div style="font-size:42px;font-weight:800;color:{color}">{final:.2f}</div>'
            f'<div style="font-size:12px;color:#64748b">/ 1.0 &nbsp; {label}</div>'
            f'</div>',
            unsafe_allow_html=True,
        )
    with pc:
        for idx, sc in enumerate(scores):
            bar_color = "#22c55e" if sc >= 0.7 else ("#f59e0b" if sc >= 0.5 else "#ef4444")
            st.markdown(f"**리뷰 #{idx+1}** &nbsp; `{sc:.2f}`")
            st.progress(sc)
    st.divider()

# ─── 4. Live log (collapsed when done) ────────────────────────
if st.session_state.logs:
    with st.expander("📋 전체 실시간 로그", expanded=st.session_state.running):
        st.code("\n".join(st.session_state.logs[-150:]), language=None)

# ─── Auto-refresh ─────────────────────────────────────────────
if st.session_state.running:
    time.sleep(0.8)
    st.rerun()

# ─── Error ────────────────────────────────────────────────────
if st.session_state.error:
    st.error(f"❌ 오류 발생\n\n{st.session_state.error}")
    st.stop()

# ─── 5. Result ────────────────────────────────────────────────
if st.session_state.done and st.session_state.result_md:
    md    = st.session_state.result_md
    title = st.session_state.topic_val
    ts    = datetime.now().strftime("%Y%m%d_%H%M%S")

    st.success(f"✅ 연구 완료!  |  `{Path(st.session_state.output_path).name}`")

    # Download buttons
    st.subheader("⬇️ 다운로드")
    d1, d2, d3 = st.columns(3)

    with d1:
        st.download_button(
            "📄 Markdown 다운로드",
            data=md.encode("utf-8"),
            file_name=f"research_{ts}.md",
            mime="text/markdown",
            use_container_width=True,
        )
    with d2:
        try:
            st.download_button(
                "📝 Word (.docx) 다운로드",
                data=_to_docx(md, title),
                file_name=f"research_{ts}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
        except Exception as e:
            st.warning(f"Word 생성 실패: {e}")
    with d3:
        try:
            st.download_button(
                "🖨️ PDF 다운로드",
                data=_to_pdf(md, title),
                file_name=f"research_{ts}.pdf",
                mime="application/pdf",
                use_container_width=True,
            )
        except Exception as e:
            st.warning(f"PDF 생성 실패: {e}")

    st.divider()

    # Result rendering
    st.subheader("📄 연구 결과")
    tab1, tab2 = st.tabs(["📖 렌더링 보기", "📋 Markdown 원문"])
    with tab1:
        st.markdown(md)
    with tab2:
        st.text_area("원문", md, height=600, label_visibility="collapsed")
