# utils/export.py
# Word(.docx) / PDF 고품질 출력 — IEEE 논문 / 컨설팅 리포트 템플릿

import io
import os
import re
from datetime import datetime

# ── 모드별 색상 테마 ───────────────────────────────────────────
_THEMES = {
    "academic": dict(
        primary=(20,  79, 168),   # IEEE 딥블루
        light=(232, 240, 255),    # 연청색 배경
        cover_bg=(10,  22,  65),  # 커버 네이비
        cover_accent=(90, 155, 255),
        label="IEEE Academic Research Paper",
        hex_primary="144FA8",
        hex_light="E8F0FF",
    ),
    "strategy": dict(
        primary=(18,  88,  55),   # 포레스트그린
        light=(228, 248, 238),    # 연두색 배경
        cover_bg=(10,  42,  28),  # 커버 다크그린
        cover_accent=(70, 190, 120),
        label="AI Strategy Report",
        hex_primary="125837",
        hex_light="E4F8EE",
    ),
}

_KR_FONT_PATHS = [
    r"C:\Windows\Fonts\malgun.ttf",
    r"C:\Windows\Fonts\NanumGothic.ttf",
]


# ══════════════════════════════════════════════════════════════
# 공통 헬퍼
# ══════════════════════════════════════════════════════════════

def _strip_inline(text: str) -> str:
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"\*(.*?)\*",     r"\1", text)
    text = re.sub(r"`([^`]+)`",     r"\1", text)
    return text.strip()


def _wrap_title(title: str, max_chars: int = 40) -> list[str]:
    words = title.split()
    lines, cur = [], ""
    for w in words:
        if cur and len(cur) + 1 + len(w) > max_chars:
            lines.append(cur)
            cur = w
        else:
            cur = (cur + " " + w).strip()
    if cur:
        lines.append(cur)
    return lines or [title]


def _find_kr_font() -> str | None:
    for p in _KR_FONT_PATHS:
        if os.path.exists(p):
            return p
    return None


# ══════════════════════════════════════════════════════════════
# LaTeX 수식 → 일반 텍스트 전처리
# ($...$, $$...$$ 를 Word/PDF 변환 전에 읽을 수 있는 평문으로 치환)
# ══════════════════════════════════════════════════════════════

# 내용만 남기고 명령어 자체는 버리는 래퍼 명령어 (\text{X} -> X)
_MATH_UNWRAP_CMDS = [
    "text", "mathrm", "mathcal", "mathbb", "mathbf", "mathsf", "mathnormal",
    "boldsymbol", "textbf", "textit", "textrm", "emph", "operatorname",
]

# 인자를 감싸는 형태로 바꾸는 명령어 (\frac{a}{b} -> (a/b))
_MATH_FUNC_CMDS = {
    "overline":  "avg({})",
    "bar":       "avg({})",
    "hat":       "hat({})",
    "tilde":     "approx({})",
    "dot":       "d/dt({})",
    "vec":       "vec({})",
    "sqrt":      "sqrt({})",
}

# 기호/그리스 문자 명령어 (백슬래시 제외한 이름 -> 유니코드)
_MATH_SYMBOLS = {
    "alpha": "α", "beta": "β", "gamma": "γ", "delta": "δ",
    "epsilon": "ε", "varepsilon": "ε", "zeta": "ζ", "eta": "η",
    "theta": "θ", "vartheta": "θ", "iota": "ι", "kappa": "κ",
    "lambda": "λ", "mu": "μ", "nu": "ν", "xi": "ξ", "pi": "π",
    "rho": "ρ", "sigma": "σ", "tau": "τ", "upsilon": "υ",
    "phi": "φ", "varphi": "φ", "chi": "χ", "psi": "ψ", "omega": "ω",
    "Gamma": "Γ", "Delta": "Δ", "Theta": "Θ", "Lambda": "Λ",
    "Xi": "Ξ", "Pi": "Π", "Sigma": "Σ", "Upsilon": "Υ",
    "Phi": "Φ", "Psi": "Ψ", "Omega": "Ω",
    "cdot": "·", "times": "×", "pm": "±", "mp": "∓", "div": "÷",
    "leq": "≤", "le": "≤", "geq": "≥", "ge": "≥",
    "neq": "≠", "ne": "≠", "approx": "≈", "sim": "~",
    "equiv": "≡", "propto": "∝",
    "in": "∈", "notin": "∉", "subset": "⊂", "subseteq": "⊆",
    "supset": "⊃", "supseteq": "⊇", "cup": "∪", "cap": "∩",
    "emptyset": "∅", "forall": "∀", "exists": "∃",
    "infty": "∞", "partial": "∂", "nabla": "∇",
    "rightarrow": "→", "to": "→", "leftarrow": "←", "gets": "←",
    "Rightarrow": "⇒", "Leftarrow": "⇐", "leftrightarrow": "↔",
    "langle": "<", "rangle": ">",
    "ldots": "…", "cdots": "⋯", "dots": "…",
    "sum": "Σ", "prod": "Π", "int": "∫",
    "top": "⊤", "bot": "⊥", "perp": "⊥",
}


def _apply_scripts(expr: str) -> str:
    """X^abc, X_abc 형태를 'X^(abc)', 'X_(abc)' 괄호 표기로 통일.
    PDF 코어 폰트(Helvetica 등)가 유니코드 위/아래첨자 글자를
    지원하지 않아 렌더링 시 깨질 위험이 있으므로, 폰트에 안전한
    ASCII 괄호 표기만 사용한다."""

    def _sup_repl(m):
        return f"^({m.group(1)})"

    def _sub_repl(m):
        return f"_({m.group(1)})"

    expr = re.sub(r"\^([A-Za-z0-9+\-=]+)", _sup_repl, expr)
    expr = re.sub(r"_([A-Za-z0-9+\-=]+)", _sub_repl, expr)
    return expr


def _brace_arg(s: str, i: int) -> tuple[str, int]:
    """s[i] == '{' 라고 가정하고 (중괄호 안 내용, 닫는 중괄호 다음 위치)를 반환"""
    depth = 0
    for j in range(i, len(s)):
        if s[j] == "{":
            depth += 1
        elif s[j] == "}":
            depth -= 1
            if depth == 0:
                return s[i + 1:j], j + 1
    return "", i + 1


def _unwrap_command(expr: str, cmd: str) -> str:
    """\\cmd{X} -> X"""
    pattern = re.compile(r"\\" + cmd + r"\{")
    while True:
        m = pattern.search(expr)
        if not m:
            return expr
        content, end = _brace_arg(expr, m.end() - 1)
        expr = expr[:m.start()] + content + expr[end:]


def _transform_command(expr: str, cmd: str, fmt: str) -> str:
    """\\cmd{A}{B}... -> fmt.format(A, B, ...) (fmt 안의 '{}' 개수만큼 인자를 소비)"""
    n_args = fmt.count("{}")
    pattern = re.compile(r"\\" + cmd + r"\{")
    while True:
        m = pattern.search(expr)
        if not m:
            return expr
        args, pos, ok = [], m.end() - 1, True
        for _ in range(n_args):
            if pos >= len(expr) or expr[pos] != "{":
                ok = False
                break
            content, pos = _brace_arg(expr, pos)
            args.append(content)
        if not ok:
            expr = expr[:m.start()] + expr[m.end():]
            continue
        expr = expr[:m.start()] + fmt.format(*args) + expr[pos:]


def _transform_underbrace(expr: str) -> str:
    """\\underbrace{A}_{B} -> A (B)"""
    pattern = re.compile(r"\\underbrace\{")
    while True:
        m = pattern.search(expr)
        if not m:
            return expr
        content, pos = _brace_arg(expr, m.end() - 1)
        sub = ""
        if expr[pos:pos + 2] == "_{":
            sub, pos = _brace_arg(expr, pos + 1)
        replacement = f"{content} ({sub})" if sub else content
        expr = expr[:m.start()] + replacement + expr[pos:]


def _latex_to_text(expr: str) -> str:
    """LaTeX 수식 내용(구분자 $ 제외)을 사람이 읽을 수 있는 평문으로 변환"""
    expr = expr.strip()

    for _ in range(6):
        prev = expr
        expr = _transform_command(expr, "frac", "({}/{})")
        expr = _transform_underbrace(expr)
        for cmd, fmt in _MATH_FUNC_CMDS.items():
            expr = _transform_command(expr, cmd, fmt)
        for cmd in _MATH_UNWRAP_CMDS:
            expr = _unwrap_command(expr, cmd)
        if expr == prev:
            break

    # \left( \right] 등 - 구분자 문자만 남기고 명령어 제거
    expr = re.sub(r"\\left\s*", "", expr)
    expr = re.sub(r"\\right\s*", "", expr)

    # 이스케이프된 특수문자 / 간격 명령
    expr = expr.replace(r"\{", "{").replace(r"\}", "}")
    expr = expr.replace(r"\%", "%").replace(r"\&", "&").replace(r"\_", "_")
    expr = re.sub(r"\\[,;:!]", " ", expr)
    expr = expr.replace(r"\ ", " ")

    # 기호 / 그리스 문자 (긴 이름이 짧은 이름의 접두어여도 단어 경계로 안전하게 매칭)
    for name, sym in _MATH_SYMBOLS.items():
        expr = re.sub(r"\\" + name + r"(?![a-zA-Z])", sym, expr)

    # 아래/위첨자 중괄호 제거: X_{abc} -> X_abc, X^{abc} -> X^abc
    expr = re.sub(r"_\{([^{}]*)\}", r"_\1", expr)
    expr = re.sub(r"\^\{([^{}]*)\}", r"^\1", expr)

    # 남은 미지의 명령어: \foo -> foo (안전망)
    expr = re.sub(r"\\([a-zA-Z]+)", r"\1", expr)

    expr = expr.replace("{", "").replace("}", "")

    # 위/아래첨자를 유니코드로 변환 (불가능하면 괄호 표기로 폴백)
    expr = _apply_scripts(expr)

    expr = re.sub(r"\s+", " ", expr).strip()
    return expr


def preprocess_math(md_text: str) -> str:
    """마크다운 안의 LaTeX 수식($...$, $$...$$)을 일반 텍스트로 치환.

    - 블록 수식 $$...$$  →  '*[수식] ...*' 형태의 독립된 줄
    - 인라인 수식 $...$  →  괄호/기호를 풀어쓴 평문으로 치환 후 문장에 그대로 삽입
    코드 블록(```...```) 안의 내용은 건드리지 않는다.
    """
    def _block_repl(m):
        return f"\n\n*[수식] {_latex_to_text(m.group(1))}*\n\n"

    def _inline_repl(m):
        return _latex_to_text(m.group(1))

    def _convert(segment: str) -> str:
        segment = re.sub(r"\$\$(.+?)\$\$", _block_repl, segment, flags=re.S)
        segment = re.sub(r"\$([^\n$]+?)\$", _inline_repl, segment)
        return segment

    parts = re.split(r"(```.*?```)", md_text, flags=re.S)
    return "".join(p if p.startswith("```") else _convert(p) for p in parts)


# ══════════════════════════════════════════════════════════════
# ASCII 아트 박스 제거 (Figure PNG와 중복되는 텍스트 다이어그램)
# ══════════════════════════════════════════════════════════════

_BOX_DRAW_CHARS = set("┌┐└┘├┤┬┴┼─│╔╗╚╝╠╣╦╩╬═║▲▼◄►")

# 언어 태그가 없는(=파이썬 등 실제 코드가 아닌) 펜스 블록만 대상으로 함
_FENCE_RE = re.compile(r"```[ \t]*\n(.*?)```\n*", re.S)
_CAPTION_RE = re.compile(r"\*Figure\s+\d+[:.][^\n]*\*\n*")


def remove_ascii_art_blocks(md_text: str) -> str:
    """박스 드로잉 문자로 그려진 ASCII 아트 다이어그램을
    'OO는 Figure N 참조' 텍스트로 대체 (실제 Figure PNG가 뒤에 삽입되므로 중복 제거).
    언어 태그가 있는 코드 펜스(```python 등)는 건드리지 않는다.
    """
    result, pos = [], 0
    for m in _FENCE_RE.finditer(md_text):
        code = m.group(1)
        result.append(md_text[pos:m.start()])
        if any(ch in _BOX_DRAW_CHARS for ch in code):
            cap_m = _CAPTION_RE.match(md_text, m.end())
            fig_num_m = (re.search(r"Figure\s+(\d+)", cap_m.group(0)) if cap_m else None) \
                or re.search(r"Figure\s+(\d+)", md_text[max(0, m.start() - 300):m.start()])
            ref = f"시스템 아키텍처는 Figure {fig_num_m.group(1)} 참조" if fig_num_m else "아래 그림 참조"
            result.append(ref + "\n\n")
            pos = cap_m.end() if cap_m else m.end()
        else:
            result.append(m.group(0))
            pos = m.end()
    result.append(md_text[pos:])
    return "".join(result)


# ══════════════════════════════════════════════════════════════
# DOCX 헬퍼
# ══════════════════════════════════════════════════════════════

def _xml_shading(para, fill_hex: str):
    """단락 배경 색상 (XML)"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)
    pPr.append(shd)


def _xml_left_border(para, color_hex: str, sz: str = "18"):
    """단락 왼쪽 테두리 (XML) — Abstract 강조용"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"),   "single")
    left.set(qn("w:sz"),    sz)
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), color_hex)
    pBdr.append(left)
    pPr.append(pBdr)


def _xml_hr(doc):
    """수평선 단락 (XML bottom border)"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Pt
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "4")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "BBBBBB")
    pBdr.append(bot)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)


def _inline_runs(para, text: str, size_pt=None, color_rgb=None):
    """**bold** / *italic* / `code` 인라인 마크다운 → Run"""
    from docx.shared import Pt
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            run = para.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`") and len(part) > 2:
            run = para.add_run(part[1:-1])
            run.font.name = "Courier New"
            if size_pt:
                run.font.size = Pt(size_pt - 0.5)
        else:
            run = para.add_run(_strip_inline(part))
        if size_pt:
            run.font.size = Pt(size_pt)
        if color_rgb:
            run.font.color.rgb = color_rgb


def _docx_heading(doc, text: str, level: int, color_rgb, size_pt: float,
                  space_before: float = 8, space_after: float = 3,
                  uppercase: bool = False):
    from docx.shared import Pt
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    content = text.upper() if uppercase else text
    run = p.add_run(content)
    run.font.size      = Pt(size_pt)
    run.font.bold      = True
    run.font.color.rgb = color_rgb
    return p


def _docx_code_block(doc, code_text: str):
    from docx.shared import Pt
    lines = code_text.split("\n") if code_text else [""]
    for ln in lines:
        p = doc.add_paragraph()
        _xml_shading(p, "F3F4F6")
        p.paragraph_format.space_before = Pt(0.5)
        p.paragraph_format.space_after  = Pt(0.5)
        from docx.shared import Cm
        p.paragraph_format.left_indent  = Cm(0.3)
        run = p.add_run(ln if ln else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(8.5)


# ══════════════════════════════════════════════════════════════
# 테이블 / 그림 헬퍼
# ══════════════════════════════════════════════════════════════

def _is_table_line(line: str) -> bool:
    return bool(re.match(r"^\s*\|", line.rstrip()))


def _parse_md_table(lines: list) -> tuple:
    if len(lines) < 3:
        return [], []
    headers = [_strip_inline(h) for h in lines[0].split("|") if h.strip()]
    rows = []
    for line in lines[2:]:
        cells = [_strip_inline(c) for c in line.split("|") if c.strip()]
        if cells:
            rows.append(cells)
    return headers, rows


def _xml_shading_cell(cell, fill_hex: str):
    """표 셀 배경색 (XML)"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _xml_cell_border(cell, color_hex: str = "BBBBBB"):
    """셀 4면 명시적 테두리 (XML) — 일관된 렌더링 보장"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"),   "single")
        b.set(qn("w:sz"),    "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color_hex)
        tcBorders.append(b)
    tcPr.append(tcBorders)


def _docx_table(doc, headers: list, rows: list, color_rgb, pr_hex: str, lgt_hex: str):
    """마크다운 테이블 → Word 표 (헤더 색상·줄무늬·명시적 테두리)"""
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    n_cols = max(len(headers), max((len(r) for r in rows), default=1))
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = "Table Grid"

    # 헤더 행
    hdr = table.rows[0]
    for i in range(n_cols):
        cell = hdr.cells[i]
        h_text = headers[i] if i < len(headers) else ""
        cell.text = h_text
        _xml_shading_cell(cell, pr_hex)
        _xml_cell_border(cell, "FFFFFF")
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(255, 255, 255)

    # 데이터 행
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        fill = lgt_hex if ri % 2 == 0 else "FFFFFF"
        for ci in range(n_cols):
            cell = row.cells[ci]
            val = row_data[ci] if ci < len(row_data) else ""
            cell.text = val
            if ri % 2 == 0:
                _xml_shading_cell(cell, fill)
            _xml_cell_border(cell, "CCCCCC")
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)

    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(2)
    sp.paragraph_format.space_after  = Pt(4)


def _docx_figure(doc, png_bytes: bytes, caption: str):
    """PNG bytes → Word 이미지 + 캡션"""
    import io as _io
    from docx.shared import Cm, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    try:
        buf = _io.BytesIO(png_bytes)
        doc.add_picture(buf, width=Cm(13.5))
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception:
        return

    cap_p = doc.add_paragraph()
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_p.paragraph_format.space_before = Pt(2)
    cap_p.paragraph_format.space_after  = Pt(8)
    run = cap_p.add_run(caption)
    run.italic = True
    run.font.size = Pt(9)


def _pdf_cell_safe(pdf, w, h, txt, border=0, fill=False, align="L"):
    """cell() 렌더링 — 폰트가 지원 못하는 문자여도 배경색·테두리는 항상 유지"""
    try:
        pdf.cell(w, h, txt, border=border, fill=fill, align=align)
        return
    except Exception:
        pass
    try:
        safe_txt = txt.encode("latin-1", errors="replace").decode("latin-1")
        pdf.cell(w, h, safe_txt, border=border, fill=fill, align=align)
    except Exception:
        pdf.cell(w, h, "", border=border, fill=fill, align=align)


def _pdf_wrap_text(pdf, text: str, max_width: float) -> list:
    """단어 단위 줄바꿈: text를 max_width(mm, 현재 폰트 기준) 안에 들어가는 줄 리스트로 변환"""
    text = text if text else ""
    words = text.split(" ")
    lines, cur = [], ""
    for w in words:
        trial = (cur + " " + w).strip() if cur else w
        if not cur or pdf.get_string_width(trial) <= max_width:
            cur = trial
        else:
            lines.append(cur)
            cur = w
    if cur or not lines:
        lines.append(cur)
    return lines


def _pdf_table(pdf, headers: list, rows: list, R: int, G: int, B: int,
               font_name: str = "Helvetica",
               lgt: tuple = (232, 240, 255)):
    """마크다운 테이블 → PDF 표 (테마 색상 적용, 컬럼 폭에 맞춰 자동 줄바꿈)"""
    if not headers or not rows:
        return
    n_cols  = max(len(headers), max((len(r) for r in rows), default=1))
    avail   = pdf.w - pdf.l_margin - pdf.r_margin
    col_w   = avail / n_cols
    x0      = pdf.l_margin
    c_pad   = 2.0
    text_w  = max(col_w - 2 * c_pad, 5)
    line_h  = 4.2
    header_style = "B" if font_name == "Helvetica" else ""

    def _draw_row(cells_text, fill_rgb, text_rgb, style, size):
        pdf.set_font(font_name, style=style, size=size)
        wrapped = [_pdf_wrap_text(pdf, t, text_w) for t in cells_text]
        n_lines = max(len(w) for w in wrapped)
        row_h = n_lines * line_h + 2

        # 페이지 하단을 넘어가면 새 페이지로
        if pdf.get_y() + row_h > pdf.h - pdf.b_margin:
            pdf.add_page()

        y0 = pdf.get_y()
        pdf.set_fill_color(*fill_rgb)
        pdf.set_text_color(*text_rgb)
        pdf.set_draw_color(204, 204, 204)
        for ci in range(n_cols):
            x = x0 + ci * col_w
            pdf.rect(x, y0, col_w, row_h, style="DF")
            lines = wrapped[ci] if ci < len(wrapped) else [""]
            top = y0 + (row_h - len(lines) * line_h) / 2
            for li, line in enumerate(lines):
                pdf.set_xy(x, top + li * line_h)
                _pdf_cell_safe(pdf, col_w, line_h, line, border=0, fill=False, align="C")
        pdf.set_xy(x0, y0 + row_h)

    # 헤더 행
    header_cells = [headers[i] if i < len(headers) else "" for i in range(n_cols)]
    _draw_row(header_cells, (R, G, B), (255, 255, 255), header_style, 8.5)

    # 데이터 행
    for ri, row_data in enumerate(rows):
        row_cells = [row_data[ci] if ci < len(row_data) else "" for ci in range(n_cols)]
        fill_rgb = lgt if ri % 2 == 0 else (255, 255, 255)
        _draw_row(row_cells, fill_rgb, (0, 0, 0), "", 8)

    pdf.ln(3)
    pdf.set_fill_color(255, 255, 255)
    pdf.set_text_color(0, 0, 0)
    pdf.set_draw_color(0, 0, 0)


def _pdf_figure(pdf, png_bytes: bytes, caption: str, font_name: str = "Helvetica"):
    """PNG bytes → PDF 이미지 + 캡션"""
    import io as _io
    avail = pdf.w - pdf.l_margin - pdf.r_margin
    img_w = min(avail, 130)
    x_pos = pdf.l_margin + (avail - img_w) / 2

    try:
        pdf.image(_io.BytesIO(png_bytes), x=x_pos, w=img_w)
    except Exception:
        return

    pdf.ln(2)
    pdf.set_font(font_name, size=8)
    pdf.set_text_color(100, 100, 100)
    pdf.set_x(pdf.l_margin)
    avail_c = pdf.w - pdf.l_margin - pdf.r_margin
    try:
        pdf.multi_cell(avail_c, 5, caption, align="C")
    except Exception:
        pass
    pdf.set_text_color(0, 0, 0)
    pdf.ln(4)


# ══════════════════════════════════════════════════════════════
# DOCX 메인
# ══════════════════════════════════════════════════════════════

def to_docx(md_text: str, title: str, mode: str = "academic",
            figures: list = None) -> bytes:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    md_text = remove_ascii_art_blocks(md_text)
    md_text = preprocess_math(md_text)

    theme    = _THEMES.get(mode, _THEMES["academic"])
    pr_color = RGBColor(*theme["primary"])
    lgt_hex  = theme["hex_light"]
    pr_hex   = theme["hex_primary"]

    doc = Document()

    # ── 페이지 설정 ──────────────────────────────────────────
    sec = doc.sections[0]
    sec.page_width    = Cm(21.0)
    sec.page_height   = Cm(29.7)
    sec.top_margin    = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin   = Cm(3.0) if mode == "strategy" else Cm(2.5)
    sec.right_margin  = Cm(2.5)

    # ── 기본 폰트 ────────────────────────────────────────────
    doc.styles["Normal"].font.name = "맑은 고딕"
    doc.styles["Normal"].font.size = Pt(10)

    # ── 표제 블록 ─────────────────────────────────────────────
    # 제목
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after  = Pt(4)
    rt = p_title.add_run(title)
    rt.font.size      = Pt(16)
    rt.font.bold      = True
    rt.font.color.rgb = pr_color

    # 템플릿 라벨
    p_lbl = doc.add_paragraph()
    p_lbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_lbl.paragraph_format.space_after = Pt(2)
    rl = p_lbl.add_run(theme["label"])
    rl.font.size       = Pt(9)
    rl.font.italic     = True
    rl.font.color.rgb  = RGBColor(100, 100, 100)

    # 생성일 / 시스템명
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_date.paragraph_format.space_after = Pt(6)
    rd = p_date.add_run(
        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}  |  "
        "Multi-Agent Research System"
    )
    rd.font.size      = Pt(8)
    rd.font.color.rgb = RGBColor(140, 140, 140)

    _xml_hr(doc)

    # ── 본문 파싱 ────────────────────────────────────────────
    is_abstract    = False
    is_refs        = False
    in_figures_sec = False
    in_code        = False
    code_buf:  list[str] = []
    table_buf: list[str] = []

    # 인라인 Figures 섹션이 있으면 appendix 대신 인라인 삽입
    has_inline_figs = bool(figures) and "## Figures" in md_text
    figures_queue   = list(figures) if has_inline_figs else []

    lines_iter = md_text.split("\n")
    line_idx   = 0

    while line_idx < len(lines_iter):
        line = lines_iter[line_idx]
        line_idx += 1

        # 코드 블록
        if line.startswith("```"):
            if table_buf:
                headers, rows = _parse_md_table(table_buf)
                if headers and rows:
                    _docx_table(doc, headers, rows, pr_color, pr_hex, lgt_hex)
                table_buf.clear()
            if in_code:
                _docx_code_block(doc, "\n".join(code_buf))
                code_buf.clear()
            in_code = not in_code
            continue
        if in_code:
            code_buf.append(line)
            continue

        # 마크다운 테이블 수집
        if _is_table_line(line):
            table_buf.append(line)
            continue
        else:
            if table_buf:
                headers, rows = _parse_md_table(table_buf)
                if headers and rows:
                    _docx_table(doc, headers, rows, pr_color, pr_hex, lgt_hex)
                table_buf.clear()

        stripped = line.strip()

        # H1
        if line.startswith("# ") and not line.startswith("## "):
            txt = _strip_inline(line[2:])
            if txt == title or txt == _strip_inline(title):
                continue
            is_abstract    = False
            in_figures_sec = False
            is_refs = any(k in txt.lower() for k in ("references", "참고"))
            _docx_heading(doc, txt, 1, pr_color, 13,
                          uppercase=(mode == "academic"))
            continue

        # H2
        if line.startswith("## "):
            txt = _strip_inline(line[3:])
            is_abstract    = any(k in txt.lower() for k in ("abstract", "초록"))
            is_refs        = any(k in txt.lower() for k in ("references", "참고"))
            in_figures_sec = "figures" in txt.lower()
            _docx_heading(doc, txt, 2, pr_color, 11,
                          space_before=8, space_after=3)
            continue

        # H3 — Figures 섹션 내부이면 PNG 삽입
        if line.startswith("### "):
            is_abstract = False
            txt = _strip_inline(line[4:])
            _docx_heading(doc, txt, 3, pr_color, 10.5,
                          space_before=5, space_after=2)
            if in_figures_sec and has_inline_figs and figures_queue:
                fig = figures_queue.pop(0)
                if fig.get("png_bytes"):
                    _docx_figure(doc, fig["png_bytes"], fig.get("caption", ""))
            continue

        # Figures 섹션 내 캡션·경로 줄 스킵 (이미 _docx_figure에서 캡션 삽입됨)
        if in_figures_sec and has_inline_figs:
            if stripped.startswith("_") and stripped.endswith("_"):
                continue
            if stripped.startswith("[Saved:"):
                continue

        # H4
        if line.startswith("#### "):
            is_abstract = False
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(1)
            _inline_runs(p, line[5:].strip(), size_pt=10)
            if p.runs:
                p.runs[0].bold = True
            continue

        # 수평선
        if re.match(r"^-{3,}$", stripped):
            is_abstract = False
            _xml_hr(doc)
            continue

        # 빈 줄
        if not stripped:
            if not is_refs:
                sp = doc.add_paragraph()
                sp.paragraph_format.space_before = Pt(0)
                sp.paragraph_format.space_after  = Pt(1)
            continue

        # 본문
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(3)
        p.paragraph_format.alignment    = 3  # JUSTIFY

        if is_abstract:
            _xml_shading(p, lgt_hex)
            _xml_left_border(p, pr_hex, sz="12")
            p.paragraph_format.left_indent = Cm(0.3)
            _inline_runs(p, line, size_pt=9.5)

        elif is_refs and re.match(r"^\[\d+\]|\d+\.\s", stripped):
            p.paragraph_format.left_indent        = Cm(0.6)
            p.paragraph_format.first_line_indent  = Cm(-0.6)
            run = p.add_run(stripped)
            run.font.size = Pt(9)

        else:
            _inline_runs(p, line)

    # 마지막 테이블 플러시
    if table_buf:
        headers, rows = _parse_md_table(table_buf)
        if headers and rows:
            _docx_table(doc, headers, rows, pr_color, pr_hex, lgt_hex)

    # 그림 appendix — 인라인 섹션이 없는 경우에만 (구버전 호환)
    if figures and not has_inline_figs:
        _docx_heading(doc, "Figures", 2, pr_color, 11, space_before=12, space_after=4)
        for fig in figures:
            if fig.get("png_bytes"):
                cap = f"{fig.get('title', '')}. {fig.get('caption', '')}"
                _docx_figure(doc, fig["png_bytes"], cap)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ══════════════════════════════════════════════════════════════
# PDF 헬퍼
# ══════════════════════════════════════════════════════════════

def _pdf_safe(pdf, txt: str, line_h: float = 6):
    pdf.set_x(pdf.l_margin)  # cursor를 항상 좌측 마진에서 시작
    avail = pdf.w - pdf.l_margin - pdf.r_margin
    try:
        pdf.multi_cell(avail, line_h, txt)
    except Exception:
        try:
            pdf.multi_cell(avail, line_h,
                           txt.encode("latin-1", errors="replace").decode("latin-1"))
        except Exception:
            pass


def _pdf_code(pdf, code_text: str, font_name: str = "Helvetica"):
    pdf.set_fill_color(243, 244, 246)
    pdf.ln(1)
    avail = pdf.w - pdf.l_margin - pdf.r_margin
    for ln in (code_text.split("\n") or [""]):
        pdf.set_font(font_name, size=8)  # 한국어 지원 폰트 사용
        pdf.set_x(pdf.l_margin)
        try:
            pdf.multi_cell(avail, 4.5, ln or " ", fill=True)
        except Exception:
            try:
                pdf.multi_cell(avail, 4.5,
                               (ln or " ").encode("latin-1", errors="replace").decode("latin-1"),
                               fill=True)
            except Exception:
                pass
    pdf.ln(2)
    pdf.set_fill_color(255, 255, 255)


# ══════════════════════════════════════════════════════════════
# PDF 메인
# ══════════════════════════════════════════════════════════════

def to_pdf(md_text: str, title: str, mode: str = "academic",
           figures: list = None) -> bytes:
    from fpdf import FPDF

    md_text = remove_ascii_art_blocks(md_text)
    md_text = preprocess_math(md_text)

    theme = _THEMES.get(mode, _THEMES["academic"])
    R, G, B = theme["primary"]
    kr_path  = _find_kr_font()

    # ── PDF 클래스 (헤더/푸터 포함) ────────────────────────
    class ResearchPDF(FPDF):
        def __init__(self, th, fn, short_t):
            super().__init__(orientation="P", unit="mm", format="A4")
            self._th      = th
            self._fn      = fn
            self._short_t = short_t

        def header(self):
            if self.page_no() == 1:
                return
            r, g, b = self._th["primary"]
            self.set_font(self._fn, size=8)
            self.set_text_color(r, g, b)
            self.set_x(12)
            short = self._short_t
            self.cell(0, 5, short, align="L")
            self.ln(1)
            self.set_draw_color(r, g, b)
            self.line(12, self.get_y(), 198, self.get_y())
            self.ln(4)
            self.set_text_color(0, 0, 0)

        def footer(self):
            if self.page_no() == 1:
                return
            r, g, b = self._th["primary"]
            self.set_y(-15)
            self.set_draw_color(r, g, b)
            self.line(12, self.get_y(), 198, self.get_y())
            self.ln(2)
            self.set_font(self._fn, size=8)
            self.set_text_color(120, 120, 120)
            self.cell(
                0, 5,
                f"Multi-Agent Research System  |  {datetime.now().strftime('%Y-%m-%d')}",
                align="L",
            )
            self.set_x(12)
            self.cell(0, 5, f"Page {self.page_no()} / {{nb}}", align="R")
            self.set_text_color(0, 0, 0)

    short_t = (title[:52] + "…") if len(title) > 52 else title
    pdf = ResearchPDF(theme, "Helvetica", short_t)
    pdf.alias_nb_pages("{nb}")

    # 마진·자동줄바꿈은 첫 페이지 전에 설정
    pdf.set_margins(14, 14, 14)
    pdf.set_auto_page_break(auto=True, margin=22)

    font_name = "Helvetica"
    if kr_path:
        try:
            pdf.add_font("KR", fname=kr_path)
            font_name = "KR"
            pdf._fn   = "KR"
        except Exception:
            pass

    def sf(size=10, bold=False):
        st = "B" if bold and font_name == "Helvetica" else ""
        pdf.set_font(font_name, style=st, size=size)

    # ── 표지 페이지 ─────────────────────────────────────────
    pdf.add_page()
    pdf.set_auto_page_break(auto=False)

    # 배경
    cbg = theme["cover_bg"]
    pdf.set_fill_color(*cbg)
    pdf.rect(0, 0, 210, 297, "F")

    # 상단 악센트 바
    acc = theme["cover_accent"]
    pdf.set_fill_color(*acc)
    pdf.rect(0, 0, 210, 5, "F")

    # 하단 악센트 바
    pdf.set_fill_color(*acc)
    pdf.rect(0, 280, 210, 5, "F")

    # 모드 라벨
    sf(10)
    pdf.set_text_color(*acc)
    pdf.set_y(60)
    pdf.set_x(0)
    pdf.cell(0, 8, theme["label"].upper(), align="C")

    # 구분선
    pdf.ln(4)
    pdf.set_draw_color(*acc)
    pdf.line(40, pdf.get_y(), 170, pdf.get_y())

    # 제목
    title_lines = _wrap_title(title, max_chars=38)
    n = len(title_lines)
    start_y = max(95, 120 - n * 10)
    pdf.set_y(start_y)
    sf(20 if n == 1 else 17, bold=True)
    pdf.set_text_color(235, 242, 255)
    for ln in title_lines:
        pdf.set_x(0)
        pdf.cell(0, 13, ln, align="C")
        pdf.ln(13)

    # 생성 정보
    pdf.set_y(230)
    sf(9)
    pdf.set_text_color(170, 190, 220)
    pdf.set_x(0)
    pdf.cell(0, 7, datetime.now().strftime("%Y년 %m월 %d일"), align="C")
    pdf.ln(7)
    pdf.set_x(0)
    pdf.cell(0, 7, "Multi-Agent Research System", align="C")

    # ── 본문 페이지 ──────────────────────────────────────────
    pdf.set_auto_page_break(auto=True, margin=22)
    pdf.add_page()
    pdf.set_x(pdf.l_margin)  # 커버 페이지 cursor 영향 초기화

    lgt = theme.get("light", (232, 240, 255))

    in_code        = False
    in_figures_sec = False
    code_buf:  list[str] = []
    table_buf: list[str] = []
    is_refs        = False

    has_inline_figs = bool(figures) and "## Figures" in md_text
    figures_queue   = list(figures) if has_inline_figs else []

    for line in md_text.split("\n"):
        # 코드 펜스
        if line.startswith("```"):
            if table_buf:
                headers, rows = _parse_md_table(table_buf)
                if headers and rows:
                    _pdf_table(pdf, headers, rows, R, G, B, font_name, lgt=lgt)
                table_buf.clear()
            if in_code:
                _pdf_code(pdf, "\n".join(code_buf), font_name)
                code_buf.clear()
            in_code = not in_code
            continue
        if in_code:
            code_buf.append(line)
            continue

        # 마크다운 테이블
        if _is_table_line(line):
            table_buf.append(line)
            continue
        else:
            if table_buf:
                headers, rows = _parse_md_table(table_buf)
                if headers and rows:
                    _pdf_table(pdf, headers, rows, R, G, B, font_name, lgt=lgt)
                table_buf.clear()

        clean = _strip_inline(line)

        # H1
        if line.startswith("# ") and not line.startswith("## "):
            txt = _strip_inline(line[2:])
            if txt == title or txt == _strip_inline(title):
                continue
            is_refs        = any(k in txt.lower() for k in ("references", "참고"))
            in_figures_sec = False
            pdf.ln(5)
            sf(13, bold=True)
            pdf.set_text_color(R, G, B)
            display = txt.upper() if mode == "academic" else txt
            _pdf_safe(pdf, display)
            y = pdf.get_y()
            pdf.set_draw_color(R, G, B)
            pdf.set_line_width(0.5)
            pdf.line(14, y, 196, y)
            pdf.set_line_width(0.2)
            pdf.ln(3)
            pdf.set_text_color(0, 0, 0)

        # H2
        elif line.startswith("## "):
            txt = _strip_inline(line[3:])
            is_refs        = any(k in txt.lower() for k in ("references", "참고"))
            in_figures_sec = "figures" in txt.lower()
            pdf.ln(4)
            sf(11, bold=True)
            pdf.set_text_color(R, G, B)
            _pdf_safe(pdf, txt)
            y = pdf.get_y()
            pdf.set_draw_color(*[min(255, c + 60) for c in (R, G, B)])
            pdf.line(14, y, 196, y)
            pdf.ln(3)
            pdf.set_text_color(0, 0, 0)

        # H3 — Figures 섹션이면 PNG 삽입
        elif line.startswith("### "):
            pdf.ln(3)
            sf(10, bold=True)
            pdf.set_text_color(R, G, B)
            _pdf_safe(pdf, _strip_inline(line[4:]))
            pdf.set_text_color(0, 0, 0)
            if in_figures_sec and has_inline_figs and figures_queue:
                fig = figures_queue.pop(0)
                if fig.get("png_bytes"):
                    _pdf_figure(pdf, fig["png_bytes"], fig.get("caption", ""), font_name)

        # H4
        elif line.startswith("#### "):
            pdf.ln(2)
            sf(10, bold=True)
            _pdf_safe(pdf, _strip_inline(line[5:]))

        # 수평선
        elif re.match(r"^-{3,}$", line.strip()):
            pdf.ln(2)
            pdf.set_draw_color(180, 180, 180)
            pdf.line(14, pdf.get_y(), 196, pdf.get_y())
            pdf.ln(4)

        # Figures 섹션 내 캡션·경로 줄 스킵
        elif in_figures_sec and has_inline_figs:
            stripped_ln = line.strip()
            if (stripped_ln.startswith("_") and stripped_ln.endswith("_")) or \
               stripped_ln.startswith("[Saved:"):
                continue
            elif clean:
                sf(10)
                _pdf_safe(pdf, clean, line_h=6)
            else:
                pdf.ln(3)

        # 본문
        elif clean:
            if is_refs:
                sf(8.5)
                indent = pdf.l_margin + 4
                pdf.set_x(indent)
                avail_ref = pdf.w - indent - pdf.r_margin
                try:
                    pdf.multi_cell(avail_ref, 5.5, clean)
                except Exception:
                    try:
                        pdf.multi_cell(avail_ref, 5.5,
                                       clean.encode("latin-1", errors="replace").decode("latin-1"))
                    except Exception:
                        pass
                pdf.ln(1)
            else:
                sf(10)
                _pdf_safe(pdf, clean, line_h=6)

        # 빈 줄
        else:
            pdf.ln(3)

    # 마지막 테이블 플러시
    if table_buf:
        headers, rows = _parse_md_table(table_buf)
        if headers and rows:
            _pdf_table(pdf, headers, rows, R, G, B, font_name, lgt=lgt)

    # 그림 appendix — 인라인 섹션이 없는 경우에만 (구버전 호환)
    if figures and not has_inline_figs:
        pdf.ln(5)
        sf(11, bold=True)
        pdf.set_text_color(R, G, B)
        _pdf_safe(pdf, "Figures")
        y = pdf.get_y()
        pdf.set_draw_color(R, G, B)
        pdf.line(14, y, 196, y)
        pdf.ln(4)
        pdf.set_text_color(0, 0, 0)
        for fig in figures:
            if fig.get("png_bytes"):
                sf(10)
                cap = f"{fig.get('title', '')}. {fig.get('caption', '')}"
                _pdf_figure(pdf, fig["png_bytes"], cap, font_name)

    return bytes(pdf.output())
